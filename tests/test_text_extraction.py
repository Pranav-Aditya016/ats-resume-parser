import tempfile
import unittest
from pathlib import Path

from docx import Document

from ats_parser.markdown_store import body_of, markdown_path, read_markdown, write_markdown
from ats_parser.text_extraction import (
    DocumentExtractor,
    DocxExtractor,
    ExtractedText,
    create_extractor,
)


class StubExtractor:
    """Stands in for a hosted OCR model so tests never make a network call."""

    def __init__(self, markdown="# OCR output\n\nSome text."):
        self.markdown = markdown
        self.calls = []
        self.last_usage = {"prompt_tokens": 10, "completion_tokens": 5, "total_tokens": 15}
        self.last_cost_usd = 0.001

    def extract(self, file_path: Path) -> ExtractedText:
        self.calls.append(file_path)
        return ExtractedText(self.markdown, "stub-ocr", "stub-model", 2)


class DocxExtractorTests(unittest.TestCase):
    def _docx(self, build) -> Path:
        document = Document()
        build(document)
        path = Path(tempfile.mkdtemp()) / "resume.docx"
        document.save(path)
        return path

    def test_headings_bullets_and_body_become_markdown(self):
        def build(document):
            document.add_heading("Work Experience", level=2)
            document.add_paragraph("Planning Engineer at ABC Developers")
            document.add_paragraph("MSP scheduling", style="List Bullet")

        extracted = DocxExtractor().extract(self._docx(build))

        self.assertIn("## Work Experience", extracted.markdown)
        self.assertIn("Planning Engineer at ABC Developers", extracted.markdown)
        self.assertIn("- MSP scheduling", extracted.markdown)
        self.assertEqual(extracted.extractor, "docx-local")

    def test_tables_become_markdown_tables(self):
        def build(document):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Company"
            table.cell(0, 1).text = "Duration"
            table.cell(1, 0).text = "ABC Developers"
            table.cell(1, 1).text = "5 Years"

        markdown = DocxExtractor().extract(self._docx(build)).markdown

        self.assertIn("| Company | Duration |", markdown)
        self.assertIn("| --- | --- |", markdown)
        self.assertIn("| ABC Developers | 5 Years |", markdown)

    def test_document_order_is_preserved_across_paragraphs_and_tables(self):
        def build(document):
            document.add_paragraph("Before the table")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Inside the table"
            document.add_paragraph("After the table")

        markdown = DocxExtractor().extract(self._docx(build)).markdown

        self.assertLess(markdown.index("Before the table"), markdown.index("Inside the table"))
        self.assertLess(markdown.index("Inside the table"), markdown.index("After the table"))

    def test_empty_document_is_an_error_rather_than_empty_markdown(self):
        path = self._docx(lambda document: None)
        with self.assertRaises(RuntimeError):
            DocxExtractor().extract(path)


class DocumentExtractorRoutingTests(unittest.TestCase):
    def test_docx_is_extracted_locally_and_never_sent_to_ocr(self):
        stub = StubExtractor()
        document = Document()
        document.add_paragraph("Local text only")
        path = Path(tempfile.mkdtemp()) / "resume.docx"
        document.save(path)

        extracted = DocumentExtractor(stub).extract(path)

        self.assertEqual(stub.calls, [])
        self.assertIn("Local text only", extracted.markdown)

    def test_pdf_is_routed_to_the_ocr_extractor(self):
        stub = StubExtractor()
        path = Path(tempfile.mkdtemp()) / "resume.pdf"
        path.write_bytes(b"%PDF-1.4")

        extracted = DocumentExtractor(stub).extract(path)

        self.assertEqual(stub.calls, [path])
        self.assertEqual(extracted.extractor, "stub-ocr")

    def test_unsupported_extension_is_rejected(self):
        path = Path(tempfile.mkdtemp()) / "resume.txt"
        path.write_text("hello", encoding="utf-8")
        with self.assertRaises(ValueError):
            DocumentExtractor(StubExtractor()).extract(path)

    def test_usage_is_surfaced_from_the_wrapped_ocr_extractor(self):
        extractor = DocumentExtractor(StubExtractor())
        self.assertEqual(extractor.last_usage["total_tokens"], 15)
        self.assertEqual(extractor.last_cost_usd, 0.001)

    def test_unknown_provider_is_rejected(self):
        with self.assertRaises(ValueError):
            create_extractor("nonexistent")


class MarkdownStoreTests(unittest.TestCase):
    def setUp(self):
        self.output_dir = Path(tempfile.mkdtemp())
        self.source = Path("resumes") / "A. Arul Jagan.pdf"
        self.path = markdown_path(self.output_dir, self.source)

    def test_written_document_carries_front_matter_and_body(self):
        write_markdown(self.path, self.source, ExtractedText("# Resume\n\nBody", "stub-ocr", "m1", 3))
        text = self.path.read_text(encoding="utf-8")

        self.assertTrue(text.startswith("---\n"))
        self.assertIn('source_file: "A. Arul Jagan.pdf"', text)
        self.assertIn('extractor: "stub-ocr"', text)
        self.assertIn("page_count: 3", text)
        self.assertIn("extracted_at:", text)
        self.assertEqual(read_markdown(self.path), "# Resume\n\nBody")

    def test_markdown_lands_in_a_markdown_subdirectory_named_for_the_source(self):
        self.assertEqual(self.path.parent.name, "markdown")
        self.assertEqual(self.path.name, "A. Arul Jagan.md")

    def test_null_metadata_is_written_as_valid_yaml(self):
        write_markdown(self.path, self.source, ExtractedText("Body", "docx-local"))
        text = self.path.read_text(encoding="utf-8")

        self.assertIn("model: null", text)
        self.assertIn("page_count: null", text)

    def test_body_separator_inside_the_body_is_not_treated_as_front_matter(self):
        write_markdown(self.path, self.source, ExtractedText("Intro\n\n---\n\nOutro", "stub-ocr"))
        self.assertEqual(read_markdown(self.path), "Intro\n\n---\n\nOutro")

    def test_missing_document_reads_as_none(self):
        self.assertIsNone(read_markdown(self.output_dir / "markdown" / "absent.md"))

    def test_document_without_front_matter_is_still_readable(self):
        self.assertEqual(body_of("# Just markdown\n"), "# Just markdown")

    def test_write_is_atomic_and_leaves_no_temporary_file(self):
        write_markdown(self.path, self.source, ExtractedText("Body", "stub-ocr"))
        leftovers = list(self.path.parent.glob("*.tmp"))
        self.assertEqual(leftovers, [])


if __name__ == "__main__":
    unittest.main()
